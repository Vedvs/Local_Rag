from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document

from config.settings import SUPPORTED_EXTENSIONS


@dataclass(frozen=True)
class LoadedDocument:
    source_path: Path
    relative_path: str
    text: str


def iter_source_files(data_root: Path) -> list[Path]:
    return sorted(
        path
        for path in data_root.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def _load_pdf(path: Path) -> str:
    pages: list[str] = []
    with fitz.open(path) as document:
        for page in document:
            page_text = page.get_text("text").strip()
            if page_text:
                pages.append(page_text)
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_document(path: Path, data_root: Path) -> LoadedDocument | None:
    if path.suffix.lower() == ".pdf":
        text = _load_pdf(path)
    elif path.suffix.lower() == ".docx":
        text = _load_docx(path)
    else:
        text = _load_text(path)

    cleaned = text.strip()
    if not cleaned:
        return None

    return LoadedDocument(
        source_path=path,
        relative_path=str(path.relative_to(data_root)),
        text=cleaned,
    )


def load_documents(data_root: Path) -> list[LoadedDocument]:
    documents: list[LoadedDocument] = []
    for source_path in iter_source_files(data_root):
        document = load_document(source_path, data_root)
        if document is not None:
            documents.append(document)
    return documents