from pathlib import Path

from docx import Document

from loaders.base import clean_text, metadata_for
from models.document import SourceDocument


def load_docx_documents(path: Path, data_root: Path) -> list[SourceDocument]:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    cleaned = clean_text("\n".join(paragraphs))
    if not cleaned:
        return []

    meta = metadata_for(path, data_root)
    return [SourceDocument(text=cleaned, **meta)]